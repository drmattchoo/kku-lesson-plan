import docx

from app.template_binder import render_lesson_plan
from tests.fixtures.dummy_lesson_plan_context import DUMMY_CONTEXT


def test_render_proof_produces_filled_docx(tmp_path):
    output_path = tmp_path / "render_proof.docx"
    render_lesson_plan(DUMMY_CONTEXT, output_path)

    assert output_path.exists()
    doc = docx.Document(output_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    full_text += "\n".join(c.text for t in doc.tables for r in t.rows for c in r.cells)

    assert "Human Physiology" in full_text
    assert "Autonomic Nervous System" in full_text
    assert "PLO4 :" in full_text
    assert "CLO2 :" in full_text
    assert "Sympathetic vs Parasympathetic" in full_text
    assert "60" in full_text
    # objective must be a CLO-tied action statement, distinct from the title
    assert "วิเคราะห์ความแตกต่างระหว่าง sympathetic และ parasympathetic ได้" in full_text
    # content is a numbered subtopic list, distinct from title and objective
    assert "Neurotransmitter และ receptor ของแต่ละระบบ" in full_text
    # ผลการเรียนรู้ column is computed PLO/CLO pairs (PLO4/CLO1), not bare CLO ids
    assert "PLO4/CLO1" in full_text
    assert "PLO4/CLO2" in full_text


def test_render_proof_normalizes_already_prefixed_plo_clo_ids(tmp_path):
    # real LLM extraction sometimes returns ids already prefixed ("CLO1"/"PLO4")
    # instead of bare ("1"/"4") — must never double up into "CLOCLO1"/"PLOPLO4"
    context = {
        **DUMMY_CONTEXT,
        "PLOs": [{"id": "PLO4", "text": "ตัวอย่าง PLO"}],
        "CLOs": [{"id": "CLO1", "text": "ตัวอย่าง CLO", "ploRefs": ["PLO4"]}],
        "keyPoints": [
            {**DUMMY_CONTEXT["keyPoints"][0], "cloRefs": ["CLO1"]},
        ],
    }
    output_path = tmp_path / "prefixed_ids.docx"
    render_lesson_plan(context, output_path)

    doc = docx.Document(output_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    full_text += "\n".join(c.text for t in doc.tables for r in t.rows for c in r.cells)

    assert "PLO4 :" in full_text
    assert "CLO1 :" in full_text
    assert "PLO4/CLO1" in full_text
    assert "PLOPLO" not in full_text
    assert "CLOCLO" not in full_text
