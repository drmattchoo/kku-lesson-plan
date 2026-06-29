"""One-off: rewrite templates/kku_lesson_plan.docx in place, replacing the blank
KKU template's dotted placeholders with Jinja tags docxtpl can render.
Run once against the pristine blank template; the tagged file is what gets committed
and reused by app.template_binder from then on.
"""
from pathlib import Path

import docx

REPO_ROOT = Path(__file__).resolve().parents[2]
TEMPLATE_PATH = REPO_ROOT / "templates" / "kku_lesson_plan.docx"


def set_paragraph_text(paragraph, text: str) -> None:
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return
    runs[0].text = text
    for run in runs[1:]:
        run.text = ""


def remove_paragraph(paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def tag_paragraphs(doc: docx.document.Document) -> None:
    p = doc.paragraphs
    set_paragraph_text(p[2], "ภาคการศึกษา {{ semester }} ปีการศึกษา {{ academicYear }}")
    set_paragraph_text(p[3], "รหัสวิชาและชื่อวิชา  {{ courseCode }} {{ courseName }}")
    set_paragraph_text(p[4], "หัวข้อเรื่อง  {{ lectureTopic }}        จำนวน {{ totalDurationHours }} ชั่วโมง")
    set_paragraph_text(p[5], "ผู้เรียน  {{ learners }}")
    set_paragraph_text(p[6], "อาจารย์ผู้สอน  {{ instructorName }}      สังกัด สาขาวิชา {{ department }}")
    set_paragraph_text(p[7], "วันที่ {{ sessionDate }}   เวลา {{ sessionTime }}")

    # PLO/CLO block: each {%p for/endfor%} tag must be ALONE in its own paragraph
    # (docxtpl repeats whatever paragraphs sit between the open and close tags).
    p[10].insert_paragraph_before("{%p for plo in PLOs %}", style=p[10].style)
    set_paragraph_text(p[10], "PLO {{ plo.id }} : {{ plo.text }}")
    p[11].insert_paragraph_before("{%p endfor %}", style=p[10].style)
    p[11].insert_paragraph_before("{%p for clo in CLOs %}", style=p[10].style)
    set_paragraph_text(p[11], "CLO {{ clo.id }} : {{ clo.text }}")
    p[12].insert_paragraph_before("{%p endfor %}", style=p[10].style)
    remove_paragraph(p[12])

    p[15].insert_paragraph_before("{%p for ref in references %}", style=p[15].style)
    set_paragraph_text(p[15], "{{ ref }}")
    p[16].insert_paragraph_before("{%p endfor %}", style=p[15].style)
    remove_paragraph(p[16])
    remove_paragraph(p[17])
    remove_paragraph(p[18])


def tag_table(doc: docx.document.Document) -> None:
    table = doc.tables[0]
    data_row = table.rows[1]
    cells = data_row.cells
    set_paragraph_text(cells[0].paragraphs[0], "{{ kp.timeLabel }}")
    set_paragraph_text(cells[1].paragraphs[0], "{{ kp.title }}")
    set_paragraph_text(cells[2].paragraphs[0], "{{ kp.durationMin }}")
    set_paragraph_text(cells[3].paragraphs[0], "{{ kp.cloRefsText }}")
    set_paragraph_text(cells[4].paragraphs[0], "{{ kp.objective }}")
    set_paragraph_text(cells[5].paragraphs[0], "{{ kp.content }}")
    set_paragraph_text(cells[6].paragraphs[0], "{{ kp.teachingMethod }}")
    set_paragraph_text(cells[7].paragraphs[0], "{{ kp.materials }}")
    set_paragraph_text(cells[8].paragraphs[0], "{{ kp.assessment }}")

    # {%tr for/endfor%} must each be ALONE in their own row (mirrors the {%p%}
    # paragraph convention) — docxtpl repeats whatever row(s) sit in between.
    for_row = table.add_row()
    data_row._element.addprevious(for_row._element)
    set_paragraph_text(for_row.cells[0].paragraphs[0], "{%tr for kp in keyPoints %}")

    endfor_row = table.add_row()
    data_row._element.addnext(endfor_row._element)
    set_paragraph_text(endfor_row.cells[0].paragraphs[0], "{%tr endfor %}")

    # rows now: [0]=header, [1]=for-tag, [2]=templated data row, [3]=endfor-tag,
    # [4..-2]=blank example rows to drop, [-1]=total row.
    for row in table.rows[4:-1]:
        row._element.getparent().remove(row._element)

    total_row = table.rows[-1]
    set_paragraph_text(total_row.cells[2].paragraphs[0], "{{ totalDurationMin }}")


def main() -> None:
    doc = docx.Document(TEMPLATE_PATH)
    tag_paragraphs(doc)
    tag_table(doc)
    doc.save(TEMPLATE_PATH)
    print(f"Tagged {TEMPLATE_PATH}")


if __name__ == "__main__":
    main()
